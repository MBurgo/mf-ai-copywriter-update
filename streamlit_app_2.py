# ✍️ Motley Fool AI Copywriter — Pro Edition v5.1 (Gemini 3 Flash Preview)
# ----------------------------------------------------------
# • Dual-Engine: Switch between OpenAI (GPT-4) and Google (Gemini 3 Flash)
# • Internal Plan (chain‑of‑thought) stage
# • JSON {plan, copy} separation with Native JSON Mode for both engines
# • Dynamic word‑count enforcement tied to dropdown
# • Dual spinners for clearer progress feedback
# • Unique keys for every button (resolves duplicate‑ID error)
# • Few‑shot “Reference Winner” exemplars for email & sales pages
# • Slider behaviour driven by external traits_config.json
# • Persistent Session State for Variants
# • Smart Markdown-to-DOCX conversion
# • Newline Sanitizer for Update/Revise mode
# ----------------------------------------------------------

import time, json, pathlib, re
from io import BytesIO
from textwrap import dedent

import streamlit as st
import google.generativeai as genai
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from google.generativeai.types import HarmCategory, HarmBlockThreshold

# ────────────────────────────────────────────────────────────
# 0.  Global toggles
# ────────────────────────────────────────────────────────────
USE_STREAMING = False     # live token stream
AUTO_QA      = True       # self‑critique & auto‑fix loop

# ---- Model & token ceiling ---------------------------------
MAX_OUTPUT_TOKENS = 4096

# ---- Length buckets (words) --------------------------------
LENGTH_RULES = {
    "📏 Short (100–200 words)":        (100, 220),
    "📐 Medium (200–500 words)":       (200, 550),
    "📖 Long (500–1500 words)":        (500, 1600),
    "📚 Extra Long (1500–3000 words)": (1500, 3200),
    "📜 Scrolling Monster (3000+ words)": (3000, None),
}

# ────────────────────────────────────────────────────────────
# 1.  Clients & Config
# ────────────────────────────────────────────────────────────

# --- OpenAI Init ---
try:
    openai_client = OpenAI(api_key=st.secrets.openai_api_key)
except Exception:
    openai_client = None # Handle gracefully if user only wants Gemini

# --- Google Gemini Init ---
GOOGLE_AVAILABLE = False
if "google_api_key" in st.secrets:
    genai.configure(api_key=st.secrets.google_api_key)
    GOOGLE_AVAILABLE = True

# ────────────────────────────────────────────────────────────
# 1A.  Load slider‑rule configuration
# ────────────────────────────────────────────────────────────
try:
    TRAIT_CFG = json.loads(pathlib.Path("traits_config.json").read_text())
except Exception as e:
    st.error(f"🚨 CRITICAL ERROR: Could not load 'traits_config.json'.\nDetails: {e}")
    st.stop()

# ────────────────────────────────────────────────────────────
# 2.  Streamlit page & CSS
# ────────────────────────────────────────────────────────────
st.set_page_config(page_title="✍️ Foolish AI Copywriter",
                   initial_sidebar_state="expanded")
st.title("✍️ Foolish AI Copywriter")

st.markdown("""
<style>
div.stButton>button { width:100%; }
h2, h3   { margin-top:1.1em; }
ul       { margin-left:1.3em; }
strong   { color:#CF7F00; }
</style>
""", unsafe_allow_html=True)

# ────────────────────────────────────────────────────────────
# 3.  Session helpers
# ────────────────────────────────────────────────────────────
def _init(**defaults):
    for k, v in defaults.items():
        st.session_state.setdefault(k, v)

_init(generated_copy="", adapted_copy="", internal_plan="", length_choice="", variants=None)

def line(label: str, value: str) -> str:
    return f"- {label}: {value}\n" if value.strip() else ""

# ────────────────────────────────────────────────────────────
# 3A.  Slider‑rule helpers
# ────────────────────────────────────────────────────────────
def trait_rules(traits: dict) -> list[str]:
    out: list[str] = []
    for name, score in traits.items():
        cfg = TRAIT_CFG.get(name)
        if not cfg: continue

        if score >= cfg["high_threshold"]:
            out.append(cfg["high_rule"])
        elif score <= cfg["low_threshold"]:
            out.append(cfg["low_rule"])
        else:
            mid_rule = cfg.get("mid_rule")
            if mid_rule: out.append(mid_rule)
    return out

def allow_exemplar(traits: dict) -> bool:
    for name, score in traits.items():
        cfg = TRAIT_CFG.get(name, {})
        if cfg.get("high_exemplar_allowed") and score >= cfg["high_threshold"]:
            return True
    return False

# ────────────────────────────────────────────────────────────
# 4.  Prompt components
# ────────────────────────────────────────────────────────────
COUNTRY_RULES = {
    "Australia":      "Use Australian English, prices in AUD, reference the ASX.",
    "United Kingdom": "Use British English, prices in GBP, reference the FTSE.",
    "Canada":         "Use Canadian English, prices in CAD, reference the TSX.",
    "United States":  "Use American English, prices in USD, reference the S&P 500.",
}

SYSTEM_PROMPT = dedent("""
You are The Motley Fool’s senior direct‑response copy chief.

• Voice: plain English, optimistic, inclusive, lightly playful but always expert.
• Draw from Ogilvy clarity, Sugarman narrative, Halbert urgency, Cialdini persuasion.
• Use **Markdown headings** (##, ###) and standard `-` bullets for lists.
• Never promise guaranteed returns; keep compliance in mind.
• The reference examples are for inspiration only — do NOT reuse phrases verbatim.
• Return ONLY the requested copy – no meta commentary, no code fences.

{country_rules}

At the very end of the piece, append this italic line (no quotes):
*Past performance is not a reliable indicator of future results.*
""").strip()

TRAIT_EXAMPLES = {
    "Urgency": [
        "This isn't a drill — once midnight hits, your chance to secure these savings is gone forever.",
        "Time’s ticking — when the clock hits zero tonight, you’re out of luck.",
        "You have exactly one shot. Miss today’s deadline, and it's gone forever."
    ],
    "Data_Richness": [
        "Last year alone, our recommendations averaged returns 220% higher than the market average.",
        "Our analysis has identified 73% higher returns than the average ASX investor over three consecutive years.",
        "More than 85% of our recommended stocks outperformed the market last fiscal year alone."
    ],
    "Social_Proof": [
        "Thousands of investors trust Motley Fool every year to transform their financial future.",
        "Australia’s leading financial experts have rated us #1 three years in a row.",
        "Join over 125,000 smart investors who rely on Motley Fool’s stock advice every month."
    ],
    "Comparative_Framing": [
        "Think back to those who seized early opportunities in the smartphone revolution.",
        "Imagine being among the first to see Netflix’s potential in 2002. That’s the kind of opportunity we’re talking about.",
        "Just like the early days of Tesla, these stocks could define your investing success for years."
    ],
    "Imagery": [
        "When that switch flips, the next phase could accelerate even faster.",
        "Think of it as a snowball rolling downhill—small at first, but soon unstoppable.",
        "Like a rocket on the launch pad, the countdown has begun and liftoff is imminent."
    ],
    "Conversational_Tone": [
        "Look — investing can feel complicated, but what if it didn't have to be?",
        "We get it—investing can seem overwhelming. But what if you had someone guiding you every step of the way?",
        "Here’s the truth: investing doesn’t have to be complicated. Let’s simplify this together."
    ],
    "FOMO": [
        "Opportunities like these pass quickly — and regret can last forever.",
        "Don’t be the one who has to tell their friends, ‘I missed out when I had the chance.’",
        "By tomorrow, your chance to act will be history. Don’t live with that regret."
    ],
    "Repetition": [
        "This offer is for today only. Today only means exactly that: today only.",
        "Act now. This offer expires tonight. Again, it expires tonight—no exceptions.",
        "This is a limited-time deal. Limited-time means exactly that: limited-time."
    ],
}

def trait_guide(traits: dict) -> str:
    out = []
    for i, (name, score) in enumerate(traits.items(), 1):
        cfg = TRAIT_CFG.get(name, {})
        high_thresh = cfg.get("high_threshold", 8)
        shots = 3 if score >= high_thresh else 2 if score >= (high_thresh - 3) else 1
        examples = " / ".join(f"“{s}”" for s in TRAIT_EXAMPLES.get(name, [])[:shots])
        out.append(f"{i}. {name.replace('_',' ')} ({score}/10) — e.g. {examples}")
    return "\n".join(out)

# --- Micro demos & Winners ----------------------------------
EMAIL_MICRO = """
### Example Email
**Subject Line:** Last chance to lock in $119 Motley Fool membership  
**Greeting:** Hi Sarah,  
**Body:** Tonight at midnight, your opportunity to save 60 % disappears. Thousands of Australians already rely on our ASX stock tips—now it’s your turn. Click before the timer hits zero and start investing smarter.  
**CTA:** Activate my membership  
**Sign‑off:** The Motley Fool Australia Team
""".strip()

SALES_MICRO = """
### Example Sales Page
## Headline  
One Day Only—Unlock the Silver Pass for $119  

### Introduction  
Imagine having two extra experts on your side every month…

### Key Benefits  
- Double the stock picks, triple the insight  
- ASX, growth & dividend coverage in one pass  
- 400,000+ Aussie investors already on board  

### Detailed Body  
Scroll down and you’ll see why the Silver Pass could be your portfolio’s inflection point. But remember—the $119 price tag vanishes at 11:59 pm tonight.  

### CTA  
**Yes! Secure My Pass Now**
""".strip()

SALES_WINNER = SALES_MICRO 
EMAIL_WINNER = EMAIL_MICRO

# --- Structural skeletons -----------------------------------
EMAIL_STRUCT = """
### Subject Line
### Greeting
### Body (benefits, urgency, proofs)
### Call‑to‑Action
### Sign‑off
""".strip()

SALES_STRUCT = """
## Headline
### Introduction
### Key Benefit Paragraphs
### Detailed Body
### Call‑to‑Action
""".strip()

# ────────────────────────────────────────────────────────────
# 5.  Prompt builder
# ────────────────────────────────────────────────────────────
def build_prompt(copy_type, copy_struct, traits, brief, length_choice, original=None):
    exemplar = EMAIL_MICRO if copy_type.startswith("📧") else SALES_MICRO
    if allow_exemplar(traits):
        exemplar += "\n\n" + (SALES_WINNER if copy_type == "📝 Sales Page" else EMAIL_WINNER)

    hard_list = trait_rules(traits)
    hard_block = "#### Hard Requirements\n" + "\n".join(hard_list) if hard_list else ""
    
    # Enhanced instruction for Updates to prevent format drift
    if original:
        edit_block = f"""
\n\n### ORIGINAL COPY TO REVISE
{original}
### INSTRUCTION:
Rewrite the copy above using the new trait requirements. 
IMPORTANT: You MUST preserve the Markdown structure (Headings, Bullets) used in the original.
"""
    else:
        edit_block = ""

    min_len, max_len = LENGTH_RULES[length_choice]
    length_block = (f"#### Length Requirement\nWrite between **{min_len} and {max_len} words**."
                    if max_len else
                    f"#### Length Requirement\nWrite **at least {min_len} words**.")

    return f"""
{trait_guide(traits)}

{exemplar}

#### Structure to Follow
{copy_struct}

{hard_block}

#### Campaign Brief
{line('Hook', brief['hook'])}{line('Details', brief['details'])}{line('Offer', f"Special {brief['offer_price']} (Retail {brief['retail_price']}), Term {brief['offer_term']}")}{line('Reports', brief['reports'])}{line('Stocks to Tease', brief['stocks_to_tease'])}{line('Quotes/News', brief['quotes_news'])}

{length_block}

{edit_block}

Please limit bullet lists to three or fewer and favour full‑sentence paragraphs elsewhere.

### END INSTRUCTIONS
""".strip()

# ────────────────────────────────────────────────────────────
# 6.  Unified LLM helper (Supports OpenAI & Gemini)
# ────────────────────────────────────────────────────────────
def run_chat(messages, engine, expect_json=False, max_tokens=MAX_OUTPUT_TOKENS):
    """
    Handles request dispatch to either OpenAI or Gemini.
    """
    
    # --- PATH A: OPENAI ---
    if engine == "OpenAI (GPT-4)":
        if not openai_client:
            st.error("OpenAI API Key missing.")
            return ""
            
        kwargs = {"max_tokens": max_tokens}
        if expect_json:
            kwargs["response_format"] = {"type": "json_object"}
        
        # Retry loop for OpenAI
        for attempt in range(3):
            try:
                resp = openai_client.chat.completions.create(
                    model=st.secrets.get("openai_model", "gpt-4-turbo"),
                    messages=messages,
                    **kwargs
                )
                return resp.choices[0].message.content.strip()
            except Exception as e:
                time.sleep(1 + attempt)
        return ""

    # --- PATH B: GOOGLE GEMINI ---
    elif engine == "Google (Gemini 3)":
        if not GOOGLE_AVAILABLE:
            st.error("Google API Key missing in secrets.toml.")
            return ""

        # 1. Extract System Instruction vs User Prompt
        sys_msg = next((m['content'] for m in messages if m['role'] == 'system'), "")
        user_prompt = "\n\n".join([m['content'] for m in messages if m['role'] != 'system'])

        # 2. Configure Model - Uses your specific requested model
        # Falls back to secrets or hardcoded string
        model_name = st.secrets.get("google_model", "gemini-3-flash-preview")
        
        # 3. Safety Settings (Important for Copywriting)
        # We need to permit "Urgency" (Hype) without triggering Harassment filters
        safety_config = {
            HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_ONLY_HIGH,
            HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_ONLY_HIGH,
            HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_ONLY_HIGH,
            HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_ONLY_HIGH,
        }

        # 4. Generation Config
        gen_config = genai.GenerationConfig(
            temperature=0.7,
            max_output_tokens=max_tokens,
            response_mime_type="application/json" if expect_json else "text/plain"
        )

        model = genai.GenerativeModel(model_name=model_name, system_instruction=sys_msg)

        # Retry loop for Gemini
        for attempt in range(3):
            try:
                response = model.generate_content(
                    user_prompt, 
                    generation_config=gen_config,
                    safety_settings=safety_config
                )
                return response.text.strip()
            except Exception as e:
                time.sleep(1 + attempt)
                if attempt == 2: st.error(f"Gemini API Error: {e}")
        return ""

# ────────────────────────────────────────────────────────────
# 7A.  AI Pair‑editor
# ────────────────────────────────────────────────────────────
def self_qa(draft, copy_type, engine):
    if not AUTO_QA: return draft

    min_len, _ = LENGTH_RULES.get(st.session_state.length_choice, (0, None))
    word_count = len(draft.split())
    if min_len and word_count < (min_len * 0.5):
        crit = f"- Draft is only {word_count} words (Target: {min_len}). Please expand significantly."
    else:
        crit = ""

    if not crit:
        # Step 1: Critique
        msgs_crit = [{"role":"system","content":"You are an obsessive editorial QA bot."},
                     {"role":"user","content":f"Check copy for: Hard requirements, Structure matches {copy_type}, Disclaimer present. Return ONLY “PASS” or bullet fixes.\n--- COPY ---\n{draft}"}]
        
        crit = run_chat(msgs_crit, engine)

    if "PASS" in crit.upper():
        return draft

    # Step 2: Fix
    msgs_fix = [{"role":"system","content":"Revise copy to address feedback."},
                {"role":"user","content":f"Apply fixes, output full revised copy ONLY.\n### FIXES\n{crit}\n### ORIGINAL\n{draft}"}]
    
    patched = run_chat(msgs_fix, engine)
    return patched.strip() if patched else draft

# ────────────────────────────────────────────────────────────
# 7B.  Variant generator helper
# ────────────────────────────────────────────────────────────
def generate_variants(base_copy: str, engine: str, n: int = 5):
    prompt = f"""
Write {n} alternative subject‑line/headline ideas AND {n} alternative CTA button labels
for the copy below, preserving tone and urgency.
Return JSON: {{ "headlines": [...], "ctas": [...] }}

--- COPY ---
{base_copy}
--- END COPY ---
"""
    msgs = [{"role":"system","content":"You are a world‑class copywriter."},
            {"role":"user","content":prompt}]
    
    resp_text = run_chat(msgs, engine, expect_json=True)
    
    # Sanitize markdown json blocks if present (Gemini sometimes adds ```json ... ```)
    clean_text = resp_text.replace("```json", "").replace("```", "").strip()
    return json.loads(clean_text)

# ────────────────────────────────────────────────────────────
# 7C.  Smart DOCX Exporter
# ────────────────────────────────────────────────────────────
def create_docx(text):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    lines = text.split('\n')
    for line in lines:
        header_match = re.match(r'^(#{2,4})\s+(.*)', line)
        if header_match:
            level = len(header_match.group(1)) - 1
            doc.add_heading(header_match.group(2), level=level)
        else:
            clean_line = line.replace('**', '') 
            if clean_line.strip():
                doc.add_paragraph(clean_line)
    
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ────────────────────────────────────────────────────────────
# 8.  UI – Generate tab
# ────────────────────────────────────────────────────────────
tab_gen, tab_adapt = st.tabs(["✍️ Generate Copy", "🌐 Adapt Copy"])

with tab_gen:
    # --- Sidebar Controls ---
    with st.sidebar:
        st.subheader("🧠 AI Engine")
        ai_engine = st.radio("Select Model", ["OpenAI (GPT-4)", "Google (Gemini 3)"])
        
        if ai_engine == "Google (Gemini 3)" and not GOOGLE_AVAILABLE:
            st.warning("⚠️ Google API Key not found. Please check secrets.toml")

        with st.expander("🎚️ Linguistic Trait Intensity", True):
            with st.form("trait_form"):
                trait_scores = {
                    "Urgency":             st.slider("Urgency & Time Sensitivity", 1, 10, 8),
                    "Data_Richness":       st.slider("Data‑Richness & Numerical Emphasis", 1, 10, 7),
                    "Social_Proof":        st.slider("Social Proof & Testimonials", 1, 10, 6),
                    "Comparative_Framing": st.slider("Comparative Framing", 1, 10, 6),
                    "Imagery":             st.slider("Imagery & Metaphors", 1, 10, 7),
                    "Conversational_Tone": st.slider("Conversational Tone", 1, 10, 8),
                    "FOMO":                st.slider("FOMO", 1, 10, 7),
                    "Repetition":          st.slider("Repetition for Emphasis", 1, 10, 5),
                }
                update_traits = st.form_submit_button("🔄 Update Copy")

    country   = st.selectbox("🌐 Target Country", list(COUNTRY_RULES))
    copy_type = st.selectbox("Copy Type", ["📧 Email", "📝 Sales Page"])
    length_choice = st.selectbox("Desired Length", list(LENGTH_RULES))
    st.session_state.length_choice = length_choice

    st.subheader("Campaign Brief")
    hook    = st.text_area("🪝 Campaign Hook")
    details = st.text_area("📦 Product / Offer Details")

    c1, c2, c3 = st.columns(3)
    offer_price  = c1.text_input("Special Offer Price")
    retail_price = c2.text_input("Retail Price")
    offer_term   = c3.text_input("Subscription Term")

    reports         = st.text_area("📑 Included Reports")
    stocks_to_tease = st.text_input("📈 Stocks to Tease (optional)")
    st.subheader("📰 Quotes or Recent News (optional)")
    quotes_news = st.text_area("Add quotes, stats, or timely news to reference")

    show_critique = st.checkbox("🧐 Show AI critique after draft", value=False)

    def brief():
        return {"country": country, "hook": hook, "details": details,
                "offer_price": offer_price, "retail_price": retail_price,
                "offer_term": offer_term, "reports": reports,
                "stocks_to_tease": stocks_to_tease, "quotes_news": quotes_news}

    copy_struct = EMAIL_STRUCT if copy_type.startswith("📧") else SALES_STRUCT

    # ────────── Core generator ────────── #
    def generate(old=None):
        if not hook.strip() or not details.strip():
            st.warning("⚠️ Please provide at least a 'Campaign Hook' and 'Product Details'.")
            return None

        prompt_core = build_prompt(copy_type, copy_struct,
                                   trait_scores, brief(), length_choice, old)
        user_instr = dedent("""
        ### TASK
        1. Create a concise INTERNAL bullet plan covering:
           • Hook & opening flow
           • Placement of proof, urgency, CTA
           2. Then write the final copy.

        Respond ONLY as valid JSON with exactly two keys:
        {
          "plan": "<the bullet outline>",
          "copy": "<the finished marketing copy>"
        }
        """).strip()

        msgs = [
            {"role":"system",
             "content": SYSTEM_PROMPT.format(country_rules=COUNTRY_RULES[country])},
            {"role":"user",
             "content": user_instr + "\n\n" + prompt_core}
        ]

        with st.spinner(f"Crafting copy with {ai_engine}…"):
            raw_json = run_chat(msgs, ai_engine, expect_json=True)

        if not raw_json: return None

        # Clean potential markdown wrappers from Gemini
        clean_json = raw_json.replace("```json", "").replace("```", "").strip()

        try:
            data = json.loads(clean_json)
        except json.JSONDecodeError:
            data = {"plan": "", "copy": clean_json}

        st.session_state.internal_plan = data["plan"].strip()

        # Sanitize Newlines
        raw_copy = data["copy"].strip()
        if "\\n" in raw_copy:
             raw_copy = raw_copy.replace("\\n", "\n")

        with st.spinner("Polishing copy…"):
            draft = self_qa(raw_copy, copy_type, ai_engine)

            if show_critique:
                msgs_critique = [
                        {"role": "system", "content": "Give concise, constructive feedback."},
                        {"role": "user", "content": f"In 3 bullets – one strength, one weakness, one improvement.\n--- COPY ---\n{draft}"}
                ]
                crit = run_chat(msgs_critique, ai_engine)
                st.info(crit)
        
        st.session_state.variants = None
        return draft

    # --- Buttons
    if st.button("✨ Generate Copy", key="gen_generate"):
        result = generate()
        if result: st.session_state.generated_copy = result

    if update_traits and st.session_state.generated_copy:
        result = generate(st.session_state.generated_copy)
        if result: st.session_state.generated_copy = result

    # --- Display & post‑gen tools
    if st.session_state.generated_copy:
        st.subheader("📝 Current Copy")
        st.markdown(st.session_state.generated_copy)

        with st.expander("🔍 Show Internal Plan (AI outline)"):
            st.markdown(st.session_state.internal_plan or "_No plan captured_")

        st.code(st.session_state.generated_copy, language="markdown")

        if st.button("🎯 Generate 5 Alt Headlines & CTAs", key="gen_variants_btn"):
            with st.spinner(f"Brainstorming variants with {ai_engine}…"):
                st.session_state.variants = generate_variants(st.session_state.generated_copy, ai_engine)

        if st.session_state.variants:
            st.subheader("📰 Headline Ideas")
            cols = st.columns(5)
            for i, text in enumerate(st.session_state.variants.get("headlines", [])):
                with cols[i]:
                    st.markdown(f"**{i+1}.** {text}")
                    st.radio(f"Vote H{i}", ["👍", "👎"], key=f"h_vote_{i}", horizontal=True, label_visibility="collapsed")

            st.subheader("🔘 CTA Button Ideas")
            cols = st.columns(5)
            for i, text in enumerate(st.session_state.variants.get("ctas", [])):
                with cols[i]:
                    st.markdown(f"**{i+1}.** {text}")
                    st.radio(f"Vote C{i}", ["👍", "👎"], key=f"c_vote_{i}", horizontal=True, label_visibility="collapsed")

        col1, col2 = st.columns(2)
        docx_file = create_docx(st.session_state.generated_copy)
        col1.download_button("📥 Download DOCX", docx_file, "mf_copy.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                               key="gen_download")

        if col2.button("🗑️ Clear", key="gen_clear"):
            st.session_state.generated_copy = ""
            st.session_state.internal_plan = ""
            st.session_state.variants = None
            st.experimental_rerun()

# ────────────────────────────────────────────────────────────
# 9.  UI – Adapt tab
# ────────────────────────────────────────────────────────────
with tab_adapt:
    st.markdown("### Paste the original copy and select a **target country**.")
    original_text = st.text_area("Original Copy", height=250)

    colA, colB = st.columns(2)
    source_c = colA.selectbox("Original Country", list(COUNTRY_RULES))
    target_c = colB.selectbox("Target Country",
                              [c for c in COUNTRY_RULES if c != source_c])

    if st.button("🌐 Adapt Copy", key="adapt_button") and original_text.strip():
        msgs = [
            {"role":"system",
             "content": SYSTEM_PROMPT.format(country_rules=COUNTRY_RULES[target_c])},
            {"role":"user",
             "content": (
                 f"Adapt the following marketing copy for a {target_c} audience.\n"
                 "Update spelling, currency, market references; preserve tone & structure.\n\n"
                 "--- ORIGINAL COPY START ---\n"
                 f"{original_text}\n"
                 "--- ORIGINAL COPY END ---\n"
                 "### END INSTRUCTIONS"
             )}
        ]
        with st.spinner("Adapting…"):
            st.session_state.adapted_copy = run_chat(msgs, ai_engine)

    if st.session_state.adapted_copy:
        st.subheader("🌐 Adapted Copy")
        st.markdown(st.session_state.adapted_copy)

        b1, b2 = st.columns(2)
        adapt_docx = create_docx(st.session_state.adapted_copy)
        b1.download_button("📥 Download DOCX", adapt_docx, "mf_adapted.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                               key="adapt_download")
                               
        if b2.button("🗑️ Clear Adapted", key="adapt_clear"):
            st.session_state.adapted_copy = ""
